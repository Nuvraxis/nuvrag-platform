import io
import re
from dataclasses import dataclass, field
from typing import Protocol

from app.core.exceptions import DocumentProcessingError
from app.models import FileType


@dataclass(slots=True)
class TextSection:
    """A contiguous run of text plus the provenance needed to cite it later."""

    content: str
    metadata: dict[str, object] = field(default_factory=dict)


class TextExtractor(Protocol):
    def extract(self, payload: bytes) -> list[TextSection]: ...


def _clean(text: str) -> str:
    # PDF and DOCX extraction leaves ragged whitespace that inflates token counts and
    # confuses the splitter's paragraph boundaries.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


class PdfExtractor:
    def extract(self, payload: bytes) -> list[TextSection]:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError

        try:
            reader = PdfReader(io.BytesIO(payload))
        except PdfReadError as exc:
            raise DocumentProcessingError(f"Unreadable PDF: {exc}", retryable=False) from exc

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise DocumentProcessingError("PDF is password protected", retryable=False) from exc

        sections: list[TextSection] = []
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                raw = page.extract_text() or ""
            except Exception:  # noqa: BLE001 - one broken page should not fail the document
                continue
            content = _clean(raw)
            if content:
                sections.append(TextSection(content=content, metadata={"page": page_number}))
        return sections


class DocxExtractor:
    def extract(self, payload: bytes) -> list[TextSection]:
        import docx
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = docx.Document(io.BytesIO(payload))
        except (PackageNotFoundError, KeyError, ValueError) as exc:
            raise DocumentProcessingError(f"Unreadable DOCX: {exc}", retryable=False) from exc

        sections: list[TextSection] = []
        heading = None
        buffer: list[str] = []

        def flush() -> None:
            body = _clean("\n".join(buffer))
            if body:
                sections.append(
                    TextSection(content=body, metadata={"section": heading} if heading else {})
                )
            buffer.clear()

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style is not None and paragraph.style.name.startswith("Heading"):
                flush()
                heading = text
                continue
            buffer.append(text)
        flush()

        for table in document.tables:
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                for row in table.rows
            ]
            body = _clean("\n".join(row for row in rows if row))
            if body:
                sections.append(TextSection(content=body, metadata={"section": "table"}))

        return sections


class MarkdownExtractor:
    _HEADING = re.compile(r"^(#{1,6})\s+(.*)$")

    def extract(self, payload: bytes) -> list[TextSection]:
        return self.sections_from_text(_decode(payload))

    def sections_from_text(self, text: str) -> list[TextSection]:
        sections: list[TextSection] = []
        heading: str | None = None
        buffer: list[str] = []

        def flush() -> None:
            body = _clean("\n".join(buffer))
            if body:
                sections.append(
                    TextSection(content=body, metadata={"section": heading} if heading else {})
                )
            buffer.clear()

        in_code_fence = False
        for line in text.splitlines():
            if line.lstrip().startswith("```"):
                in_code_fence = not in_code_fence
                buffer.append(line)
                continue
            match = None if in_code_fence else self._HEADING.match(line)
            if match:
                flush()
                heading = match.group(2).strip()
            else:
                buffer.append(line)
        flush()
        return sections


class MdxExtractor(MarkdownExtractor):
    """MDX is Markdown with JSX and ESM mixed in.

    The component markup is layout, not prose: embedding `<Callout variant="warning">` gives
    the retriever nothing to match on and dilutes the sentence it wraps. Imports, exports,
    frontmatter and JSX tags are therefore removed, while the text between tags is kept —
    that is usually the actual documentation. Fenced code blocks are left untouched, since a
    `.mdx` file explaining JSX legitimately contains it as an example.
    """

    _FRONTMATTER = re.compile(r"\A---\n.*?\n---\n", re.DOTALL)
    _ESM = re.compile(r"^\s*(?:import|export)\s")
    _JSX_COMMENT = re.compile(r"\{\s*/\*.*?\*/\s*\}", re.DOTALL)
    _JSX_EXPRESSION = re.compile(r"\{[^{}\n]*\}")
    _JSX_TAG = re.compile(r"</?[A-Za-z][\w.:-]*(?:\s[^<>]*?)?/?>")

    def extract(self, payload: bytes) -> list[TextSection]:
        return self.sections_from_text(self._strip_jsx(_decode(payload)))

    def _strip_jsx(self, text: str) -> str:
        text = self._FRONTMATTER.sub("", text)
        kept: list[str] = []
        in_code_fence = False

        for line in text.splitlines():
            if line.lstrip().startswith("```"):
                in_code_fence = not in_code_fence
                kept.append(line)
                continue
            if in_code_fence:
                kept.append(line)
                continue
            if self._ESM.match(line):
                continue

            cleaned = self._JSX_COMMENT.sub("", line)
            cleaned = self._JSX_TAG.sub("", cleaned)
            # Runs last, so an expression that was an attribute value is already gone and
            # only standalone `{count}` interpolations remain.
            cleaned = self._JSX_EXPRESSION.sub("", cleaned)
            kept.append(cleaned)

        return "\n".join(kept)


class PlainTextExtractor:
    def extract(self, payload: bytes) -> list[TextSection]:
        content = _clean(_decode(payload))
        return [TextSection(content=content)] if content else []


def _decode(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError("Unable to decode text file", retryable=False)


_EXTRACTORS: dict[FileType, TextExtractor] = {
    FileType.PDF: PdfExtractor(),
    FileType.DOCX: DocxExtractor(),
    FileType.MD: MarkdownExtractor(),
    FileType.MDX: MdxExtractor(),
    FileType.TXT: PlainTextExtractor(),
}


def extract_text(file_type: FileType, payload: bytes) -> list[TextSection]:
    extractor = _EXTRACTORS.get(file_type)
    if extractor is None:
        raise DocumentProcessingError(f"No extractor registered for {file_type}", retryable=False)
    sections = extractor.extract(payload)
    if not sections:
        raise DocumentProcessingError(
            "No extractable text found; the file may be an image-only scan",
            retryable=False,
        )
    return sections
